#!/usr/bin/env python3
"""End-to-end tests for all document skills (7 capabilities).

Run:  python3 a2a/skills/tests/test_skills.py
Deps: pip3 install python-docx mammoth openpyxl pymupdf fpdf2
"""

import os
import shutil
import sys
import traceback

TEST_DIR = "/tmp/skill_tests"
passed = 0
failed = 0
results = []


def setup():
    os.makedirs(TEST_DIR, exist_ok=True)


def teardown():
    shutil.rmtree(TEST_DIR, ignore_errors=True)


def run_test(name, fn):
    global passed, failed
    try:
        fn()
        print(f"PASS: {name}")
        results.append(("PASS", name))
        passed += 1
    except Exception as e:
        print(f"FAIL: {name}: {e}")
        traceback.print_exc()
        results.append(("FAIL", name))
        failed += 1


# ---------------------------------------------------------------------------
# 1. DOCX write
# ---------------------------------------------------------------------------
def test_docx_write():
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("Test Report", level=0)
    doc.add_paragraph("This is a test paragraph.")

    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "A"
    table.rows[0].cells[1].text = "B"
    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "2"

    path = os.path.join(TEST_DIR, "test_write.docx")
    doc.save(path)
    assert os.path.exists(path), "DOCX file was not created"
    assert os.path.getsize(path) > 0, "DOCX file is empty"


# ---------------------------------------------------------------------------
# 2. DOCX read
# ---------------------------------------------------------------------------
def test_docx_read():
    from docx import Document

    # Create
    doc = Document()
    doc.add_paragraph("Hello from docx read test")
    path = os.path.join(TEST_DIR, "test_read.docx")
    doc.save(path)

    # Read back
    doc2 = Document(path)
    texts = [p.text for p in doc2.paragraphs]
    assert "Hello from docx read test" in texts, f"Expected text not found in {texts}"


# ---------------------------------------------------------------------------
# 3. DOCX convert (to Markdown via mammoth)
# ---------------------------------------------------------------------------
def test_docx_convert():
    from docx import Document
    import mammoth

    doc = Document()
    doc.add_heading("Convert Heading", level=1)
    doc.add_paragraph("Paragraph to convert.")
    path = os.path.join(TEST_DIR, "test_convert.docx")
    doc.save(path)

    with open(path, "rb") as f:
        result = mammoth.convert_to_markdown(f)
    md = result.value
    assert "Convert Heading" in md, f"Heading not found in markdown: {md[:200]}"
    assert "Paragraph to convert" in md, f"Paragraph not found in markdown: {md[:200]}"


# ---------------------------------------------------------------------------
# 4. Spreadsheet write
# ---------------------------------------------------------------------------
def test_spreadsheet_write():
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Test"

    ws["A1"] = "Product"
    ws["B1"] = "Price"
    ws["A2"] = "Widget"
    ws["B2"] = 9.99
    ws["A3"] = "Gadget"
    ws["B3"] = 19.99
    ws["B4"] = "=SUM(B2:B3)"

    path = os.path.join(TEST_DIR, "test_write.xlsx")
    wb.save(path)
    assert os.path.exists(path), "XLSX file was not created"
    assert os.path.getsize(path) > 0, "XLSX file is empty"


# ---------------------------------------------------------------------------
# 5. Spreadsheet read
# ---------------------------------------------------------------------------
def test_spreadsheet_read():
    from openpyxl import Workbook, load_workbook

    # Create
    wb = Workbook()
    ws = wb.active
    ws["A1"] = "Name"
    ws["B1"] = "Score"
    ws["A2"] = "Alice"
    ws["B2"] = 95
    path = os.path.join(TEST_DIR, "test_read.xlsx")
    wb.save(path)

    # Read back
    wb2 = load_workbook(path)
    ws2 = wb2.active
    assert ws2["A2"].value == "Alice", f"Expected 'Alice', got {ws2['A2'].value}"
    assert ws2["B2"].value == 95, f"Expected 95, got {ws2['B2'].value}"


# ---------------------------------------------------------------------------
# 6. Spreadsheet analyze (formulas)
# ---------------------------------------------------------------------------
def test_spreadsheet_analyze():
    from openpyxl import Workbook, load_workbook

    wb = Workbook()
    ws = wb.active
    ws["A1"] = 10
    ws["A2"] = 20
    ws["A3"] = "=SUM(A1:A2)"
    ws["B1"] = 5
    ws["B2"] = 15
    ws["B3"] = "=AVERAGE(B1:B2)"
    path = os.path.join(TEST_DIR, "test_analyze.xlsx")
    wb.save(path)

    # Read back without data_only → should see formula strings
    wb2 = load_workbook(path, data_only=False)
    ws2 = wb2.active
    assert ws2["A3"].value == "=SUM(A1:A2)", f"Expected formula, got {ws2['A3'].value}"
    assert ws2["B3"].value == "=AVERAGE(B1:B2)", f"Expected formula, got {ws2['B3'].value}"


# ---------------------------------------------------------------------------
# 7. PDF extract (create with fpdf2, read with PyMuPDF)
# ---------------------------------------------------------------------------
def test_pdf_extract():
    from fpdf import FPDF
    import fitz

    # Create
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "", 14)
    pdf.cell(0, 10, "PDF extraction roundtrip test", new_x="LMARGIN", new_y="NEXT")
    path = os.path.join(TEST_DIR, "test_extract.pdf")
    pdf.output(path)

    # Extract text with PyMuPDF
    doc = fitz.open(path)
    text = ""
    for page in doc:
        text += page.get_text()
    doc.close()

    assert "PDF extraction roundtrip test" in text, f"Expected text not found in: {text[:200]}"


# ---------------------------------------------------------------------------
# 8. PDF generate (multi-page with table)
# ---------------------------------------------------------------------------
def test_pdf_generate():
    from fpdf import FPDF
    import fitz

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Page 1: title + text
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 20)
    pdf.cell(0, 12, "Multi-Page Report", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.set_font("Helvetica", "", 12)
    pdf.multi_cell(0, 7, "This report demonstrates multi-page PDF generation. " * 10)

    # Page 2: table
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, "Data Table", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 12)

    headers = ["Item", "Qty", "Price"]
    col_widths = [60, 40, 40]
    pdf.set_font("Helvetica", "B", 12)
    for h, w in zip(headers, col_widths):
        pdf.cell(w, 10, h, border=1, align="C")
    pdf.ln()

    pdf.set_font("Helvetica", "", 12)
    for item, qty, price in [("Alpha", "10", "$100"), ("Beta", "25", "$250"), ("Gamma", "5", "$50")]:
        pdf.cell(col_widths[0], 10, item, border=1)
        pdf.cell(col_widths[1], 10, qty, border=1, align="C")
        pdf.cell(col_widths[2], 10, price, border=1, align="R")
        pdf.ln()

    path = os.path.join(TEST_DIR, "test_generate.pdf")
    pdf.output(path)

    # Verify with PyMuPDF
    doc = fitz.open(path)
    assert doc.page_count >= 2, f"Expected at least 2 pages, got {doc.page_count}"
    assert os.path.getsize(path) > 500, f"PDF too small: {os.path.getsize(path)} bytes"
    doc.close()


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    setup()
    try:
        run_test("docx_write", test_docx_write)
        run_test("docx_read", test_docx_read)
        run_test("docx_convert", test_docx_convert)
        run_test("spreadsheet_write", test_spreadsheet_write)
        run_test("spreadsheet_read", test_spreadsheet_read)
        run_test("spreadsheet_analyze", test_spreadsheet_analyze)
        run_test("pdf_extract", test_pdf_extract)
        run_test("pdf_generate", test_pdf_generate)
    finally:
        teardown()

    total = passed + failed
    print(f"\n{'All' if failed == 0 else 'Only'} {passed}/{total} tests passed")
    sys.exit(0 if failed == 0 else 1)
